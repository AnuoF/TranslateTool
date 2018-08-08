#!/usr/bin/env python
# -*- encoding: utf-8 -*-
 
'''
@Author  :   Allen
 
@License :   (C) Copyright 2018, Allen's Studio
 
@Contact :   188512936@qq.com
 
@Software:   VS2017
 
@File    :   T_Docx.py
 
@Time    :   June 21,2018
 
@Desc    :   此模块实现.docx文档的翻译功能.
 
'''


from T_Base import Translate
from TranslateFunc import *
import os
import docx    # 安装指令：pip install python-docx
from Logger import *


class DocxTranslate(Translate):

    def __init__(self, fileName, fullPath):
        '''构造函数'''
        # fileName:文件名
        # fullPath:全路径

        self.fileName = fileName
        self.fullName = fullPath
        self.prepare()


    def translate(self):
        '''翻译'''

        # 获取文档对象
        doc = docx.Document(self.fullName)

        # 创建内存中的word文档对象
        new_doc = docx.Document()

        i = 0
        # 遍历每一段文本
        for para in doc.paragraphs:
            # 翻译
            content = para.text.strip()    # 去掉多余空格

            if content != '':
                ret = translate_func(content)
                trans = ret if ret else '翻译失败'
                # 写入新文件
                new_doc.add_paragraph(content)
                new_doc.add_paragraph(trans)
                i += 1
                print(i,end=' ',flush=True)
                # 保存到本地文件
                new_doc.save(self.new_fullPath)
        
        Logger().write(self.fileName + '翻译完成，新文档：' + self.new_fullPath)


    def prepare(self):
        '''准备：生成的文件名和路径'''

        # 查看要生成的文件名是否已存在，若存在，则在文件名中 + 1
        file_name = os.path.splitext(self.fileName)[0] + os.path.splitext(self.fileName)[1]
        path = self.get_path('Doc_Out',file_name)
        i = 1

        while os.path.exists(path):   # 循环，生成新的文件名
            file_name = os.path.splitext(self.fileName)[0] + str(i) + os.path.splitext(self.fileName)[1]
            path = self.get_path('Doc_Out',file_name)
            i = i + 1

        self.new_fileName = file_name
        self.new_fullPath = path


    def get_path(self,*paths):
        '''获取，组装路径'''

        path = os.path.split(os.path.realpath(__file__))[0]
        if len(paths):
            for i in range(len(paths)):
                path = os.path.join(path,paths[i])

        return path